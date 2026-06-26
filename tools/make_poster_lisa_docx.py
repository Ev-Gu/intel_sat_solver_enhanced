#!/usr/bin/env python3
"""Generate Poster_Results_Lisa.docx from completed benchmark + fuzzer results."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

HERE = Path(__file__).resolve().parent
OUT = HERE.parent / "results" / "Poster_Results_Lisa.docx"


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = "Arial"
    r.font.color.rgb = RGBColor(0x1F, 0x3D, 0x7A)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.name = "Arial"
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def heading(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.name = "Arial"
    r.font.color.rgb = RGBColor(0x1F, 0x3D, 0x7A)


def para(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(11)
    r.font.name = "Arial"


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.name = "Arial"


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    add_title(doc, "IntelTopor MaxSAT — Poster Results")
    add_subtitle(doc, "June 26, 2026  |  Bisan / Yevgeny  |  Branch: New-Basic-Corectness-Validation")

    para(doc, "")
    para(
        doc,
        "Summary of correctness fuzzing and performance benchmarking for the IntelTopor "
        "MaxSAT solver (incremental IPAMIR + batch WCNF paths).",
    )

    heading(doc, "1. Correctness — IPAMIR differential fuzzer (Fuzzer 2)")
    para(
        doc,
        "Random WCNF instances are fed to IntelTopor IPAMIR and UWrMaxSat 1.4 IPAMIR "
        "(exact reference oracle). We compare SAT/UNSAT status and objective after each solve.",
    )
    add_table(
        doc,
        ["Run", "Instances", "Bugs", "Suboptimal"],
        [
            ["Smoke test", "5", "0", "4"],
            ["Extended test", "20", "0", "12"],
        ],
    )
    bullet(doc, "Internal per-solve timeout via TOPOR_IPAMIR_TIME_LIMIT.")
    bullet(doc, "Progress logging: c timeo <time> <cost> on improving solutions.")
    bullet(
        doc,
        "No SAT/UNSAT or hard-clause correctness bugs in the extended run "
        "(after Yevgeny's fixes + timeout integration).",
    )
    bullet(
        doc,
        "Remaining gap is suboptimality under tight NUWLS budget (1s), not wrong answers.",
    )

    heading(doc, "2. Performance — ref-first benchmark (Yam method, final run)")
    para(
        doc,
        "Benchmark set: MSE22Unique.csv regression instances. "
        "Phase 1: EvalMaxSAT 2022 filters instances it solves within 60 seconds. "
        "Phase 2: IntelTopor batch and IPAMIR run only on that fair subset.",
    )
    bullet(doc, "148 instances scanned → 141 passed EvalMaxSAT filter.")
    bullet(doc, "Timeout: 60 seconds per solver invocation (not competition 7200s).")
    add_table(
        doc,
        ["Path", "Solved", "Certified optimum", "Avg time (solved)"],
        [
            ["IntelTopor batch (-M 1)", "78 / 141", "78 / 141", "1.27 s"],
            ["IntelTopor IPAMIR", "140 / 141", "99 / 141", "26.54 s"],
            ["EvalMaxSAT 2022 (filter)", "141 / 141", "—", "~0 s"],
        ],
    )
    para(
        doc,
        "Takeaway: On instances the reference solved within 60s, the IPAMIR path solves "
        "far more problems than batch (140 vs 78). Batch is faster when it succeeds.",
        bold=True,
    )

    heading(doc, "3. Suggested poster bullet points")
    bullet(
        doc,
        "Incremental MaxSAT (IPAMIR) integrated into IntelTopor with differential fuzzing "
        "vs UWrMaxSat — 0 correctness bugs in 20-instance extended fuzz run.",
    )
    bullet(
        doc,
        "Internal timeout + c timeo align IPAMIR behavior with competition-style anytime output.",
    )
    bullet(
        doc,
        "Ref-first benchmark (141 instances, 60s): IPAMIR 140/141 solved, 99/141 certified "
        "optimum vs batch 78/141.",
    )

    heading(doc, "4. Notes / limitations")
    bullet(doc, "Performance numbers use 60s timeout and MSE22 regression subset — not full MSE2022 incremental applications.")
    bullet(doc, "Official MaxSAT Evaluation uses 7200s per instance; this run follows Yam's advice for development/poster subset.")
    bullet(doc, "NUWLS parameter alignment with tt-open-wbo-inc is planned as a follow-up (not included in these numbers).")

    heading(doc, "5. Testing infrastructure")
    add_table(
        doc,
        ["Component", "Purpose"],
        [
            ["scripts/fuzz_maxsat.sh", "Fuzzer 1: batch WCNF vs EvalMaxSAT"],
            ["scripts/fuzz_maxsat_ipamir.sh", "Fuzzer 2: IPAMIR vs UWrMaxSat"],
            ["scripts/run_mse2022_bench.sh", "Ref-first performance benchmark"],
            ["scripts/run_maxsat_regression.sh", "MSE regression correctness"],
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
